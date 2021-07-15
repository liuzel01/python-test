from docx import Document

def view_docs(docx_file):
    # 打开文档1
    doc = Document(docx_file)
    # 读取每段内容
    pl = [ paragraph.text for paragraph in doc.paragraphs]
    # 输出读取到的内容
    for i in pl:
        print(i)


def view_docs_table(docx_file):
    # 打开文档1
    doc = Document(docx_file)
    # 读取每段内容
    tables = [table for table in doc.tables]
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text, end='  ')
            print()
        print('\n')


if __name__ == '__main__':
    view_docs("Python自动化办公实战课.docx")
    view_docs_table("Python自动化办公实战课.docx")