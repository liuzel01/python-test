from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import _Cell
from docx.oxml import OxmlElement


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


##标题

def add_header(text, level, align='center'):
    title_ = document.add_heading(level=level)
    if align == 'center':
        title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
    elif align == 'right':
        title_.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    # title_run.font.size = Pt(24)  # 设置标题字体大小
    title_run.font.name = 'Times New Roman'  # 设置标题西文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置标题中文字体


add_header(text='Python自动化办公实战', level=1)
add_header(text='Python基础', level=2, align='left')

document.add_paragraph('Python 是一门面向对象的高级编程语言，易学易用，是自动化办公首选的工具。')

add_header('Python玩转图片', level=2, align='left')

document.add_paragraph('图片是工作中接触较多的媒体文件了，你可能需要图片压缩，加水印，文字识别等操作')

records = (
    ('Python 基础', '00:30', '2021-08-01', ''),
    ('Python 玩转图片', '01:00', '2021-08-01', ''),
    ('Python 玩转 word', '01:00', '2021-08-01', ''),
)

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '章节'
hdr_cells[1].text = '时长'
hdr_cells[2].text = '日期'
hdr_cells[3].text = '备注'

for cell in hdr_cells:
    set_cell_border(cell,
                    top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    bottom={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    start={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    end={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    )

for chapter, time, date, note in records:
    row_cells = table.add_row().cells
    row_cells[0].text = chapter
    row_cells[1].text = time
    row_cells[2].text = date
    row_cells[3].text = note
    for cell in row_cells:
        set_cell_border(cell,
                        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                        bottom={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                        start={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                        end={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                        )
document.save('Python自动化办公实战.docx')
